from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION_START
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import re
import os

from docx.enum.table import WD_ALIGN_VERTICAL

# Visual generation configuration for Napkin AI
# Base: 1 visual per ~500 words (~4-5 visuals per average document)
# Extra: +1 visual if document exceeds 12 pages (~3600 words)
WORDS_PER_VISUAL = 500
EXTRA_VISUAL_PAGE_THRESHOLD = 12  # pages

# Helper para logo
def preparar_logo():
    # En este entorno, asumimos que el logo ya está en static/img/logo_redaxion.png
    # Si fuera necesario moverlo a /tmp, se haría aquí, pero podemos usar la ruta directa.
    pass

def insertar_logo_encabezado_derecha(doc, ruta_logo="static/img/logo_redaxion.png", ancho_pulgadas=0.9):
    try:
        if not os.path.exists(ruta_logo):
            # Fallback to old logo if new one missing
            ruta_logo = "static/img/logo.png"
            if not os.path.exists(ruta_logo):
                print(f"⚠️ Logo no encontrado: {ruta_logo}")
                return

        section = doc.sections[0]
        header = section.header

        # Tabla con ancho total obligatorio para encabezado
        table = header.add_table(rows=1, cols=2, width=Inches(7.0))
        table.allow_autofit = True
        table.autofit = True

        # Celda izquierda vacía con ancho amplio para empujar el logo
        cell_izquierda = table.cell(0, 0)
        cell_izquierda.width = Inches(6.2)
        cell_izquierda.text = ""

        # Celda derecha con el logo
        cell_derecha = table.cell(0, 1)
        paragraph = cell_derecha.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = paragraph.add_run()
        run.add_picture(ruta_logo, width=Inches(ancho_pulgadas))

        # Ajustes de espaciado y alineación
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        cell_derecha.vertical_alignment = WD_ALIGN_VERTICAL.TOP

        # Altura mínima de la fila
        table.rows[0].height = Pt(1)

        # Frase institucional centrada debajo del logo (o en el header general)
        # Nota: El usuario pidió "Frase institucional centrada debajo del logo", 
        # pero en su código la agrega al header directamente (header.add_paragraph), 
        # lo cual la pondría debajo de la tabla.
        frase = "RedaXion – Tecnología que transforma tu estudio"
        p_frase = header.add_paragraph()
        p_frase.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_frase = p_frase.add_run(frase)
        run_frase.bold = True
        run_frase.italic = True
        run_frase.font.name = 'Calibri'
        run_frase.font.size = Pt(10)
        
    except Exception as e:
        print(f"Error al insertar logo: {e}")




ESTILOS_COLOR = {
    "azul elegante": {
        "titulo": {"fondo": "4A66AC", "letra": "FFFFFF"},
        "subtitulo": {"fondo": "D8DFEF", "letra": "000000"},
        "texto": {"fondo": "FFFFFF", "letra": "000000"},
    },
    "azul pastel": {
        "titulo": {"fondo": "4F81BD", "letra": "FFFFFF"},
        "subtitulo": {"fondo": "DCE5F0", "letra": "000000"},
        "texto": {"fondo": "FFFFFF", "letra": "000000"},
    },
    "rojo elegante": {
        "titulo": {"fondo": "C10905", "letra": "FFFFFF"},
        "subtitulo": {"fondo": "FFBFBF", "letra": "000000"},
        "texto": {"fondo": "FFFFFF", "letra": "000000"},
    },
    "rojo pastel": {
        "titulo": {"fondo": "E32E91", "letra": "FFFFFF"},
        "subtitulo": {"fondo": "F9D4E8", "letra": "000000"},
        "texto": {"fondo": "FFFFFF", "letra": "000000"},
    },
    "gris elegante": {
        "titulo": {"fondo": "7F7F7F", "letra": "FFFFFF"},
        "subtitulo": {"fondo": "E5E5E5", "letra": "000000"},
        "texto": {"fondo": "FFFFFF", "letra": "000000"},
    },
    "morado pastel": {
        "titulo": {"fondo": "B553D9", "letra": "FFFFFF"},
        "subtitulo": {"fondo": "D094E6", "letra": "000000"},
        "texto": {"fondo": "FFFFFF", "letra": "000000"},
    },
    "verde pastel": {
        "titulo": {"fondo": "569F3B", "letra": "FFFFFF"},
        "subtitulo": {"fondo": "DAEFD3", "letra": "000000"},
        "texto": {"fondo": "FFFFFF", "letra": "000000"},
    }
}

# ---------------------------
# Helpers de tipografía Calibri
# ---------------------------
def _forzar_calibri_en_estilo(doc, font_name="Calibri", font_size=11):
    """Fija Calibri en el estilo Normal y en rFonts para evitar reemplazos."""
    if 'Normal' not in doc.styles:
        return
        
    style = doc.styles['Normal']
    font = style.font
    font.name = font_name
    font.size = Pt(font_size)
    
    # Ensure element exists
    if style.element is None: 
         return
         
    # Asegurar mapeos de fuentes en XML (clave para que Word no sustituya)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:cs'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)

def _set_run_calibri(run, font_name="Calibri"):
    """Asegura Calibri en un run, incluyendo rFonts."""
    run.font.name = font_name
    # Proteger los cuatro mapas
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:cs'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)

def aplicar_estilo(parrafo, tipo_bloque, color, espaciado_extra=False):
    color_normalizado = color.strip().lower()
    if color_normalizado not in ESTILOS_COLOR:
        # print(f"⚠️ Color no reconocido: '{color}'. Usando 'azul elegante' por defecto.")
        color_normalizado = "azul elegante"

    estilos_por_color = ESTILOS_COLOR[color_normalizado]
    estilos = estilos_por_color.get(tipo_bloque, {})
    fondo = estilos.get("fondo", "FFFFFF")
    letra = estilos.get("letra", "000000")

    # Color de texto
    for run in parrafo.runs:
        run.font.color.rgb = RGBColor(
            int(letra[0:2], 16),
            int(letra[2:4], 16),
            int(letra[4:6], 16)
        )

    # Fondo del párrafo (títulos y subtítulos)
    if tipo_bloque in ["titulo", "subtitulo"]:
        p = parrafo._p
        pPr = p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), fondo)
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), "auto")
        pPr.append(shd)

    # Espaciado adicional
    if espaciado_extra:
        parrafo.paragraph_format.space_after = Pt(10)


# Import formula conversion utility
try:
    from services.formula_utils import latex_to_text
except ImportError:
    def latex_to_text(text):
        return text  # Fallback if module not available


def agregar_texto_con_negrita(parrafo, texto):
    if not texto.strip():
        _set_run_calibri(parrafo.add_run(" "))
        return

    # Convert LaTeX formulas to readable text
    texto = latex_to_text(texto)

    patron = r"\*\*(.*?)\*\*"
    cursor = 0
    for match in re.finditer(patron, texto):
        start, end = match.span()
        if start > cursor:
            r1 = parrafo.add_run(texto[cursor:start])
            _set_run_calibri(r1)
        run_negrita = parrafo.add_run(match.group(1))
        run_negrita.bold = True
        _set_run_calibri(run_negrita)
        cursor = end
    if cursor < len(texto):
        r2 = parrafo.add_run(texto[cursor:])
        _set_run_calibri(r2)

def configurar_columnas(doc, tipo):
    if tipo == "doble":
        try:
             section = doc.sections[-1]
             # Check if cols element exists, if not create logic might be more complex
             # For now trusting existing logic but wrapped in try/except
             cols = section._sectPr.xpath('./w:cols')
             if cols:
                 cols[0].set(qn('w:num'), '2')
                 cols[0].set(qn('w:space'), '720')
             else:
                 # TODO: Add logic to create columns xml if missing
                 pass
             section.left_margin = Inches(0.5)
             section.right_margin = Inches(0.5)
             section.top_margin = Inches(0.75)
             section.bottom_margin = Inches(0.75)
        except Exception as e:
            print(f"Warning: Failed to configure columns: {e}")


from services.napkin_integration import generate_napkin_visual
from services.kroki_integration import generate_kroki_visual, generate_math_visual
from io import BytesIO
import time

def procesar_linea_con_formulas(doc, linea, color, espaciado_extra=False):
    if "<formula>" not in linea.lower():
        p = doc.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(10) if espaciado_extra else Pt(4)
        p.paragraph_format.line_spacing = 1.1
        agregar_texto_con_negrita(p, linea)
        aplicar_estilo(p, "texto", color)
        return

    partes = re.split(r"(<formula>.*?</formula>)", linea, flags=re.IGNORECASE)
    
    current_p = None

    def get_or_create_p():
        nonlocal current_p
        if current_p is None:
            current_p = doc.add_paragraph()
            current_p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            current_p.paragraph_format.space_after = Pt(10) if espaciado_extra else Pt(4)
            current_p.paragraph_format.line_spacing = 1.1
        return current_p

    for parte in partes:
        if parte.lower().startswith("<formula>") and parte.lower().endswith("</formula>"):
            if current_p is not None:
                aplicar_estilo(current_p, "texto", color)
                current_p = None  # Reset for next text block
            
            eq = parte[9:-10].strip()
            img_stream = generate_math_visual(eq)
            
            p_formula = doc.add_paragraph()
            p_formula.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_formula.paragraph_format.space_before = Pt(14)
            p_formula.paragraph_format.space_after = Pt(14)
            
            if img_stream:
                run = p_formula.add_run()
                run.add_picture(img_stream, width=Inches(3.0))
            else:
                agregar_texto_con_negrita(p_formula, eq)
                aplicar_estilo(p_formula, "texto", color)
            
        else:
            if parte.strip():
                p = get_or_create_p()
                agregar_texto_con_negrita(p, parte)

    if current_p is not None:
        aplicar_estilo(current_p, "texto", color)


def guardar_como_docx(texto, path_salida="/tmp/procesado.docx", color="azul oscuro", columnas="simple"):
    preparar_logo()
    doc = Document()
    _forzar_calibri_en_estilo(doc)
    insertar_logo_encabezado_derecha(doc)
    
    # Agregar título "Transcripción" al documento TCP
    titulo = doc.add_heading("Transcripción", level=1)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    aplicar_estilo(titulo, "titulo", color)
    doc.add_paragraph("")
    
    if columnas == "doble":
        configurar_columnas(doc, columnas)

    # Márgenes estrechos
    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    # ========================================================================
    # NAPKIN AI: Intelligent Visual Selection Strategy
    # ========================================================================
    # Strategy:
    # 1. Count total words in document
    # 2. Calculate number of visuals: ~1 per 800 words (min 1, max 3)
    # 3. If document > 12 pages (~3600 words), add 1 extra visual
    # 4. Distribute visuals uniformly across major sections (##)
    # 5. Generate visuals with Napkin AI (with rate limiting)
    # ========================================================================
    
    print("\n" + "="*70)
    print("📊 ANALIZANDO DOCUMENTO PARA GENERACIÓN DE VISUALES")
    print("="*70)
    
    # Parse document structure
    texto_lineas = texto.split('\n')
    total_words = sum(len(line.split()) for line in texto_lineas if line.strip())
    
    # Find all major sections (## titles)
    sections = []
    current_section = {"title": "Introducción", "start_line": 0, "content": [], "words": 0}
    
    for idx, linea in enumerate(texto_lineas):
        linea_stripped = linea.strip()
        if linea_stripped.startswith("## ") or linea_stripped.startswith("# "):
            # Save previous section
            if current_section["content"]:
                current_section["words"] = sum(len(l.split()) for l in current_section["content"])
                sections.append(current_section)
            
            # Start new section
            title = linea_stripped.replace("## ", "").replace("# ", "").strip()
            current_section = {"title": title, "start_line": idx, "content": [], "words": 0}
        else:
            current_section["content"].append(linea_stripped)
    
    # Don't forget the last section
    if current_section["content"]:
        current_section["words"] = sum(len(l.split()) for l in current_section["content"])
        sections.append(current_section)
    
    # Calculate number of visuals
    base_visuals = min(5, max(1, total_words // WORDS_PER_VISUAL))  # 1 per 500 words, min 1, max 5
    estimated_pages = total_words // 300  # Rough estimate: ~300 words per page
    extra_visual = 1 if estimated_pages > EXTRA_VISUAL_PAGE_THRESHOLD else 0
    num_visuals = base_visuals + extra_visual
    
    print(f"📝 Total de palabras: {total_words}")
    print(f"📄 Páginas estimadas: {estimated_pages}")
    print(f"🎨 Visuales a generar: {num_visuals} (base: {base_visuals}, extra: {extra_visual})")
    print(f"📚 Secciones encontradas: {len(sections)}")
    
    # Select sections for visual generation
    # Strategy: Distribute uniformly, prioritize longer sections
    visual_sections = []
    if num_visuals > 0 and len(sections) > 0:
        # Sort sections by word count (descending)
        sections_sorted = sorted(sections, key=lambda s: s["words"], reverse=True)
        
        # Select top N sections (up to num_visuals)
        visual_sections = sections_sorted[:min(num_visuals, len(sections))]
        
        print(f"\n🎯 Secciones seleccionadas para visuales:")
        for i, sec in enumerate(visual_sections, 1):
            print(f"   {i}. \"{sec['title']}\" ({sec['words']} palabras)")
    
    # Generate visuals with Napkin AI
    visuals_data = {}  # Dictionary: section_title -> BytesIO image
    
    if visual_sections:
        print(f"\n{'='*70}")
        print(f"🎨 GENERANDO VISUALES CON NAPKIN AI")
        print(f"{'='*70}\n")
        
        for idx, section in enumerate(visual_sections, 1):
            print(f"\n[{idx}/{len(visual_sections)}] Generando visual para: \"{section['title']}\"")
            
            # Prepare content for Napkin
            # Use section title + first ~200 words of content for context
            section_text = " ".join(section["content"])
            max_context = 200
            words = section_text.split()
            context = " ".join(words[:max_context])
            
            # Combine title + context for better visual generation
            napkin_input = f"{section['title']}. {context}"
            
            # Generate visual - Napkin AI con fallback a Kroki
            img_stream = generate_napkin_visual(napkin_input, language="es")
            if not img_stream:
                print(f"🔄 Activando fallback a Kroki para \"{section['title']}\"")
                img_stream = generate_kroki_visual(napkin_input, color_theme=color)
            
            if img_stream:
                visuals_data[section["title"]] = img_stream
                print(f"✅ Visual generado para \"{section['title']}\"")
            else:
                print(f"⚠️ No se pudo generar visual para \"{section['title']}\"")
            
            # Rate limiting: wait between requests (except for last one)
            if idx < len(visual_sections):
                wait_time = 0.6  # Conservative: 2 req/sec max
                print(f"⏳ Esperando {wait_time}s (rate limiting)...")
                time.sleep(wait_time)
        
        print(f"\n{'='*70}")
        print(f"✅ Generación de visuales completada: {len(visuals_data)}/{len(visual_sections)}")
        print(f"{'='*70}\n")
    
    # ========================================================================
    # Build Document with Visuals
    # ========================================================================
    
    titulo_agregado = False
    current_section_title = ""
    
    for linea in texto_lineas:
        linea = linea.rstrip()
        if not linea.strip():
            continue

        linea_normalizada = linea.lstrip()
        
        # Track current section
        if linea_normalizada.startswith("## ") or linea_normalizada.startswith("# "):
            current_section_title = linea_normalizada.replace("## ", "").replace("# ", "").strip()
        
        # Check if we should insert a visual BEFORE this section's content
        if current_section_title in visuals_data and linea_normalizada.startswith("### "):
            # Insert visual after first subsection of a major section
            # This gives good placement without interrupting the main title
            img_stream = visuals_data[current_section_title]
            
            try:
                # 10% larger images for better visibility (3.2->3.52, 5.0->5.5)
                # NOTA: Reducido otra vez a pedido del usuario (2.3 para doble, 3.8 para simple)
                doc.add_picture(img_stream, width=Inches(2.3) if columnas=="doble" else Inches(3.8))
                last_p = doc.paragraphs[-1]
                last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # Aesthetic Spacing
                last_p.paragraph_format.space_before = Pt(14)
                last_p.paragraph_format.space_after = Pt(18)
                print(f"🖼️ Visual insertado para sección: \"{current_section_title}\"")
                
                # Remove from dict so we don't insert again
                del visuals_data[current_section_title]
            except Exception as e:
                print(f"❌ Error insertando visual para \"{current_section_title}\": {e}")

        # 🔁 Tolerancia: tratar #### como ### para subtítulos mal formateados
        if linea_normalizada.startswith("#### "):
            linea_normalizada = linea_normalizada.replace("#### ", "### ")

        if linea_normalizada.startswith("### "):
            texto_subtitulo = linea_normalizada.replace("### ", "")
            p = doc.add_paragraph()
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = p.add_run(texto_subtitulo)
            run.font.size = Pt(12)
            run.font.bold = True
            _set_run_calibri(run)
            aplicar_estilo(p, "subtitulo", color)

        elif linea_normalizada.startswith("## ") or linea_normalizada.startswith("# "):
            texto_titulo = linea_normalizada.replace("## ", "").replace("# ", "")
            p = doc.add_paragraph()
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = p.add_run(texto_titulo)
            run.font.size = Pt(14)
            run.font.bold = True
            _set_run_calibri(run)
            aplicar_estilo(p, "titulo", color)
            if columnas == "doble" and not titulo_agregado:
                configurar_columnas(doc, columnas)
                titulo_agregado = True
        else:
            procesar_linea_con_formulas(doc, linea_normalizada, color)

    doc.save(path_salida)
    return path_salida






def guardar_quiz_como_docx(texto_preguntas_y_respuestas, path_guardado="/tmp/quiz.docx", color="azul oscuro", columnas="simple"):
    preparar_logo()
    doc = Document()
    _forzar_calibri_en_estilo(doc)  # <<< Fuerza Calibri 11
    insertar_logo_encabezado_derecha(doc)

    # Márgenes estrechos
    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    if columnas == "doble":
        configurar_columnas(doc, columnas)

    # (Se mantiene tu seteo explícito del estilo Normal)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Limpia encabezado "Respuestas" si viene duplicado
    texto_preguntas_y_respuestas = re.sub(
        r"(?i)^\s*\*{0,3}\s*respuestas\s*\*{0,3}[:：]?\s*$",
        "",
        texto_preguntas_y_respuestas,
        flags=re.MULTILINE
    )

    titulo = doc.add_heading("Quiz", level=1)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    aplicar_estilo(titulo, "titulo", color)
    doc.add_paragraph("")

    # Mejorada la detección de la sección de solucionario generada por GPT
    match = re.search(r"(##\s*solucionario|\*{0,2}solucionario\*{0,2}|\*{0,2}respuestas:?\*{0,2}|respuesta 1:)", texto_preguntas_y_respuestas.lower())
    if not match:
        preguntas = texto_preguntas_y_respuestas.strip()
        respuestas = ""
    else:
        idx = match.start()
        preguntas = texto_preguntas_y_respuestas[:idx].strip()
        respuestas = texto_preguntas_y_respuestas[idx:].strip()

    bloques = re.split(r"\n(?=\d{1,2}\. )", preguntas)
    for bloque in bloques:
        lineas = bloque.strip().split("\n")
        if not lineas:
            continue

        encabezado = lineas[0].strip()
        # Limpiar markdown headers (## PREGUNTAS -> PREGUNTAS)
        encabezado = re.sub(r'^#{1,6}\s*', '', encabezado)
        
        if "preguntas de práctica" in encabezado.lower() or "preguntas" == encabezado.lower():
            p = doc.add_paragraph()
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = p.add_run("Preguntas de práctica" if "práctica" in encabezado.lower() else "Preguntas")
            run.bold = True
            _set_run_calibri(run)
            run.font.size = Pt(13)
            aplicar_estilo(p, "subtitulo", color, espaciado_extra=True)
            p.paragraph_format.space_after = Pt(6)
        else:
            # Skip empty headers after cleaning
            if not encabezado:
                continue
            p = doc.add_paragraph()
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = p.add_run(encabezado)
            run.bold = True
            _set_run_calibri(run)
            run.font.size = Pt(12)
            aplicar_estilo(p, "subtitulo", color, espaciado_extra=True)

        for linea in lineas[1:]:
            linea = linea.strip()
            if not linea:
                continue
            # Limpiar markdown headers de las líneas también
            linea = re.sub(r'^#{1,6}\s*', '', linea)
            if not linea:
                continue
            # Se mantienen **...** para que agregar_texto_con_negrita aplique bold
            procesar_linea_con_formulas(doc, linea, color)

    doc.add_page_break()

    if respuestas:
        p = doc.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run("Respuestas")
        run.bold = True
        _set_run_calibri(run)
        run.font.size = Pt(14)
        aplicar_estilo(p, "titulo", color)

        for linea in respuestas.strip().split("\n"):
            linea = linea.strip()
            if not linea:
                continue
            if re.match(r"^\s*(\#\#\s*solucionario|\*{0,2}respuestas\*{0,2})\s*[:：]?\s*$", linea, flags=re.IGNORECASE):
                continue
            if "importante: esta sección es obligatoria" in linea.lower():
                continue
            # ❌ Sin limpieza de **...**
            procesar_linea_con_formulas(doc, linea, color)

    doc.save(path_guardado)
    return path_guardado


# --- Logic from convertidor_pdf.py ---

import subprocess
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.utils import simpleSplit


def hex_to_rgb(hex_str: str):
    """Convierte hex string (e.g. 'FFFFFF') a tuple floats (1.0, 1.0, 1.0)"""
    if not hex_str: return (0, 0, 0)
    try:
        r = int(hex_str[0:2], 16) / 255.0
        g = int(hex_str[2:4], 16) / 255.0
        b = int(hex_str[4:6], 16) / 255.0
        return (r, g, b)
    except:
        return (0, 0, 0)


def clean_markdown_for_pdf(text):
    """Remove markdown formatting for PDF fallback rendering."""
    import re
    # Remove ## headers (keep text)
    text = re.sub(r'^#{1,6}\s*', '', text)
    # Remove ** bold markers (keep text)
    text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
    return text.strip()

def fallback_pdf_conversion(path_docx, path_pdf, color="azul elegante"):
    """
    Fallback avanzado: Lee el DOCX e intenta replicar formato completo (Negrita, Color, Tamaño, Fondos)
    en el PDF usando ReportLab, similar al DOCX original.
    """
    try:
        print(f"⚠️ Generando PDF fallback 'Rich' para: {path_docx}")
        
        doc = Document(path_docx)
        
        # Normalizar color
        color_scheme = color.strip().lower()
        if color_scheme not in ESTILOS_COLOR:
            color_scheme = "azul elegante"
        
        print(f"🎨 Usando esquema de color: {color_scheme}")
        
        c = canvas.Canvas(path_pdf, pagesize=letter)
        width, height = letter
        
        # Margins - más compactos
        margin_left = 50
        margin_right = 50
        margin_top = 40  # Reducido de 50
        margin_bottom = 50
        
        # Start Y position
        y = height - margin_top
        
        # --- LOGO COMPACTO ---
        base_dir = os.getcwd()
        logo_path = os.path.join(base_dir, "static/img/logo_redaxion.png")
        
        if not os.path.exists(logo_path):
            logo_path = os.path.join(base_dir, "static/img/logo.png")
            
        if os.path.exists(logo_path):
             try:
                 # Logo más pequeño
                 logo_w = 0.9 * 72  # 0.9 inches
                 logo_h = 0.35 * 72 
                 # Draw top right
                 c.drawImage(logo_path, width - margin_right - logo_w, height - 30, 
                             width=logo_w, height=logo_h, preserveAspectRatio=True, mask='auto')
                 
                 # Slogan CENTRADO y más cerca del logo
                 c.setFont("Helvetica-Oblique", 8)
                 c.setFillColorRGB(0.4, 0.4, 0.4)
                 slogan = "RedaXion – Tecnología que transforma tu estudio"
                 c.drawCentredString(width / 2, height - 35, slogan)
                 
                 # Menos espacio consumido
                 y -= 45  # Reducido de 50
             except Exception as e:
                 print(f"Error pintando logo: {e}")

        max_width = width - margin_left - margin_right
        
        # Obtener estilos de color
        estilos = ESTILOS_COLOR.get(color_scheme, ESTILOS_COLOR["azul elegante"])
        
        for para in doc.paragraphs:
            if not para.text.strip():
                y -= 8  # Spacing reducido
                continue

            # Detectar tipo de párrafo
            text = para.text.strip()
            is_title = False
            is_subtitle = False
            
            # Limpiar markdown del texto (## y **)
            text = clean_markdown_for_pdf(text)
            
            # Heurística: detectar títulos por texto específico o tamaño
            if para.text.strip().startswith("#"):
                is_title = True
            elif "Quiz" in text or "Transcripción" in text or "RedaXion" in text or text.startswith("Preguntas de"):
                # Títulos especiales como "RedaQuiz", "Preguntas de práctica", etc.
                is_title = True
            elif para.runs and para.runs[0].font.size:
                size_pt = int(para.runs[0].font.size.pt) if para.runs[0].font.size.pt else 11
                if size_pt >= 14:
                    is_title = True
                elif size_pt >= 12:
                    is_subtitle = True
            
            # Defaults
            font_name = "Helvetica"
            font_size = 11
            text_color = (0, 0, 0)
            bg_color = None
            
            if para.runs and len(para.runs) > 0:
                r = para.runs[0]
                if r.bold:
                    font_name = "Helvetica-Bold"
                
                if r.font.size:
                    try:
                        font_size = int(r.font.size.pt)
                    except:
                        pass
                
                # Color del texto
                if r.font.color and r.font.color.rgb:
                    text_color = hex_to_rgb(str(r.font.color.rgb))
            
            # Aplicar estilos según tipo
            if is_title:
                font_name = "Helvetica-Bold"
                font_size = max(font_size, 14)
                # Colores del título
                bg_hex = estilos["titulo"]["fondo"]
                text_hex = estilos["titulo"]["letra"]
                bg_color = hex_to_rgb(bg_hex)
                text_color = hex_to_rgb(text_hex)
            elif is_subtitle:
                font_name = "Helvetica-Bold"
                font_size = max(font_size, 12)
                # Colores del subtítulo
                bg_hex = estilos["subtitulo"]["fondo"]
                text_hex = estilos["subtitulo"]["letra"]
                bg_color = hex_to_rgb(bg_hex)
                text_color = hex_to_rgb(text_hex)
            
            line_height = font_size * 1.3
            
            # Envolver texto
            wrapped_lines = simpleSplit(text, font_name, font_size, max_width)
            
            # Dibujar fondo si es título o subtítulo
            if bg_color and len(wrapped_lines) > 0:
                # Calcular altura del bloque
                block_height = len(wrapped_lines) * line_height + 6  # padding
                
                if y - block_height < margin_bottom:
                    c.showPage()
                    y = height - margin_top
                
                # Dibujar rectángulo de fondo
                c.setFillColorRGB(*bg_color)
                c.rect(margin_left - 5, y - block_height + line_height - 3, 
                       max_width + 10, block_height, fill=True, stroke=False)
            
            # Dibujar texto
            c.setFont(font_name, font_size)
            c.setFillColorRGB(*text_color)
            
            for line in wrapped_lines:
                if y < margin_bottom:
                    c.showPage()
                    y = height - margin_top
                    c.setFont(font_name, font_size)
                    c.setFillColorRGB(*text_color)
                
                c.drawString(margin_left, y, line)
                y -= line_height
            
            # Espaciado entre párrafos
            y -= line_height * 0.3

        c.save()
        print(f"✅ PDF fallback Rich generado con colores: {path_pdf}")
        return path_pdf

    except Exception as e:
        print(f"❌ Error en fallback PDF Rich: {e}")
        import traceback
        traceback.print_exc()
        return None




def convert_to_pdf(path_docx, color="azul elegante"):
    """
    Convierte un archivo .docx a .pdf usando LibreOffice.
    Si falla, usa fallback con ReportLab.
    Retorna la ruta al PDF generado o None si falló todo.
    """
    output_dir = os.path.dirname(path_docx)
    path_pdf = path_docx.replace(".docx", ".pdf")

    # Intentar LibreOffice
    try:
        # Check if libreoffice is installed (simplified check by running)
        subprocess.run([
            "libreoffice", "--headless", "--convert-to", "pdf", path_docx, "--outdir", output_dir
        ], check=True, capture_output=True)
        
        if os.path.exists(path_pdf):
            print(f"✅ PDF generado correctamente con LibreOffice: {path_pdf}")
            return path_pdf
    except Exception as e:
        print(f"⚠️ LibreOffice no disponible o falló: {e}")

    # Fallback si LibreOffice falló o no generó archivo
    if not os.path.exists(path_pdf):
        return fallback_pdf_conversion(path_docx, path_pdf, color)
    
    return path_pdf

