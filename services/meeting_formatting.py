"""
Meeting Formatting Service - Creates DOCX and PDF documents for meeting minutes

Produces formal meeting documents with:
- RedaXion logo header
- Structured sections (Decisions, Actions, Blockers, etc.)
- Clean, professional formatting
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os
import re

# Import utilities from main formatting module
from services.formatting import (
    insertar_logo_encabezado_derecha,
    preparar_logo,
    fallback_pdf_conversion,
    convert_to_pdf
)


def _set_run_calibri(run, font_name="Calibri"):
    """Ensure Calibri font on a run."""
    run.font.name = font_name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:cs'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)


def agregar_texto_con_negrita(parrafo, texto):
    """Add text to paragraph, converting **text** to bold."""
    if not texto.strip():
        _set_run_calibri(parrafo.add_run(" "))
        return

    patron = r"\*\*(.*?)\*\*"
    cursor = 0
    for match in re.finditer(patron, texto):
        start, end = match.span()
        if start > cursor:
            r1 = parrafo.add_run(texto[cursor:start])
            _set_run_calibri(r1)
        run_negrita = parrafo.add_run(match.group(1))
        run_negrita.bold = True
        _set_run_calibri(run_negrita)
        cursor = end
    if cursor < len(texto):
        r2 = parrafo.add_run(texto[cursor:])
        _set_run_calibri(r2)


def parse_markdown_table(lines: list) -> list:
    """
    Parse a markdown table from lines.
    Returns list of rows, each row is a list of cell values.
    """
    rows = []
    for line in lines:
        line = line.strip()
        if not line or line.startswith('|--') or line.startswith('|-'):
            continue  # Skip separator lines
        if line.startswith('|') and line.endswith('|'):
            cells = [cell.strip() for cell in line[1:-1].split('|')]
            rows.append(cells)
    return rows


def guardar_acta_reunion_como_docx(contenido: str, path_salida: str = "/tmp/acta_reunion.docx") -> str:
    """
    Save meeting minutes content to a formal DOCX document.
    
    Args:
        contenido: Markdown-like text from ChatGPT with headers, tables, and bold
        path_salida: Output path for the DOCX file
        
    Returns:
        Path to the saved DOCX file
    """
    preparar_logo()
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Add logo header
    insertar_logo_encabezado_derecha(doc)
    
    # Set margins
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    
    # Process content line by line
    lineas = contenido.split('\n')
    i = 0
    
    while i < len(lineas):
        linea = lineas[i].rstrip()
        
        # Skip empty lines but add spacing
        if not linea.strip():
            doc.add_paragraph("")
            i += 1
            continue
        
        linea_normalizada = linea.lstrip()
        
        # Handle --- separator lines
        if linea_normalizada.startswith('---'):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run("─" * 80)
            run.font.color.rgb = RGBColor(28, 198, 194)  # Accent cyan
            run.font.size = Pt(8)
            _set_run_calibri(run)
            i += 1
            continue
        
        # Main headers (## TITLE)
        if linea_normalizada.startswith('## '):
            texto = linea_normalizada.replace('## ', '')
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(texto)
            run.bold = True
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(74, 102, 172)  # Blue accent
            _set_run_calibri(run)
            i += 1
            continue
        
        # Sub headers (### subtitle)
        if linea_normalizada.startswith('### '):
            texto = linea_normalizada.replace('### ', '')
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(texto)
            run.bold = True
            run.font.size = Pt(12)
            _set_run_calibri(run)
            i += 1
            continue
        
        # Check for markdown table
        if linea_normalizada.startswith('|'):
            # Collect all table lines
            table_lines = []
            while i < len(lineas) and lineas[i].strip().startswith('|'):
                table_lines.append(lineas[i])
                i += 1
            
            # Parse and create table
            rows = parse_markdown_table(table_lines)
            if rows:
                num_cols = len(rows[0])
                table = doc.add_table(rows=len(rows), cols=num_cols)
                table.style = 'Table Grid'
                
                for row_idx, row_data in enumerate(rows):
                    for col_idx, cell_text in enumerate(row_data):
                        if col_idx < num_cols:
                            cell = table.cell(row_idx, col_idx)
                            cell.text = cell_text
                            
                            # Header row styling
                            if row_idx == 0:
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.bold = True
                                        run.font.size = Pt(10)
                            else:
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.font.size = Pt(10)
                
                doc.add_paragraph("")  # Spacing after table
            continue
        
        # Bullet points
        if linea_normalizada.startswith('- '):
            texto = linea_normalizada[2:]
            p = doc.add_paragraph(style='List Bullet')
            agregar_texto_con_negrita(p, texto)
            i += 1
            continue
        
        # Numbered items
        if re.match(r'^\d+\.', linea_normalizada):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(3)
            agregar_texto_con_negrita(p, linea_normalizada)
            i += 1
            continue
        
        # Regular paragraph
        p = doc.add_paragraph()
        agregar_texto_con_negrita(p, linea_normalizada)
        i += 1
    
    doc.save(path_salida)
    print(f"✅ Acta de reunión DOCX guardada: {path_salida}")
    return path_salida


def guardar_acta_reunion_como_pdf(contenido: str, path_pdf: str = "/tmp/acta_reunion.pdf") -> str:
    """
    Save meeting minutes to a formal PDF document.
    First creates DOCX, then converts to PDF.
    
    Args:
        contenido: Markdown-like text from ChatGPT
        path_pdf: Output path for the PDF file
        
    Returns:
        Path to the saved PDF file
    """
    # First create DOCX
    path_docx = path_pdf.replace('.pdf', '.docx')
    guardar_acta_reunion_como_docx(contenido, path_docx)
    
    # Convert to PDF
    result_pdf = convert_to_pdf(path_docx, color="azul elegante")
    
    if result_pdf and os.path.exists(result_pdf):
        print(f"✅ Acta de reunión PDF guardada: {result_pdf}")
        return result_pdf
    
    # Fallback
    print("⚠️ Usando fallback para PDF de acta")
    return fallback_pdf_conversion(path_docx, path_pdf, color="azul elegante")
