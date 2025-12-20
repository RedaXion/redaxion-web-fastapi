"""
Exam Formatting Service - Creates formal DOCX and PDF documents for exams

Produces simple, formal exam documents with:
- RedaXion logo header
- Headers with bold text
- Clean formatting suitable for printing
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os
import re

# Import PDF utilities from main formatting module
from services.formatting import (
    insertar_logo_encabezado_derecha,
    preparar_logo,
    fallback_pdf_conversion,
    convert_to_pdf
)
from services.formula_utils import latex_to_text


def _set_run_calibri(run, font_name="Calibri"):
    """Ensure Calibri font on a run."""
    run.font.name = font_name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:cs'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)


def agregar_texto_con_negrita(parrafo, texto):
    """Add text to paragraph, converting **text** to bold."""
    if not texto.strip():
        _set_run_calibri(parrafo.add_run(" "))
        return

    patron = r"\*\*(.*?)\*\*"
    cursor = 0
    for match in re.finditer(patron, texto):
        start, end = match.span()
        if start > cursor:
            r1 = parrafo.add_run(texto[cursor:start])
            _set_run_calibri(r1)
        run_negrita = parrafo.add_run(match.group(1))
        run_negrita.bold = True
        _set_run_calibri(run_negrita)
        cursor = end
    if cursor < len(texto):
        r2 = parrafo.add_run(texto[cursor:])
        _set_run_calibri(r2)


# Color dictionary for exam styling
COLORES_EXAMEN = {
    "azul elegante": RGBColor(0, 75, 135),      # Deep blue
    "verde académico": RGBColor(0, 100, 60),    # Forest green
    "rojo institucional": RGBColor(140, 20, 20), # Burgundy
    "morado moderno": RGBColor(90, 0, 120),     # Purple
    "naranja vibrante": RGBColor(180, 80, 0)    # Burnt orange
}


def obtener_color_titulo(color: str) -> RGBColor:
    """Get RGB color for exam titles/headers."""
    return COLORES_EXAMEN.get(color.lower(), COLORES_EXAMEN["azul elegante"])


def guardar_examen_como_docx(contenido: str, path_salida: str = "/tmp/examen.docx", color: str = "azul elegante") -> str:
    """
    Save exam content to a formal DOCX document.
    
    Args:
        contenido: Markdown-like text from ChatGPT with ## headers and **bold**
        path_salida: Output path for the DOCX file
        color: Color scheme for document styling
        
    Returns:
        Path to the saved DOCX file
    """
    preparar_logo()
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Add logo header
    insertar_logo_encabezado_derecha(doc)
    
    # Set margins
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    
    # Process content line by line
    lineas = contenido.split('\n')
    
    for linea in lineas:
        linea_original = linea
        linea = linea.rstrip()
        
        # Skip empty lines - don't add extra spacing
        if not linea.strip():
            continue
        
        linea_normalizada = linea.lstrip()
        
        # Convert LaTeX formulas to readable text
        linea_normalizada = latex_to_text(linea_normalizada)
        
        # Handle [dejar espacio] or similar - add blank lines for answer space
        if '[dejar espacio' in linea_normalizada.lower() or '[espacio para respuesta' in linea_normalizada.lower():
            for _ in range(8):  # Add 8 blank lines for answer space
                doc.add_paragraph("")
            continue
        
        # Handle --- separator lines - SKIP them, don't add visual separators
        if linea_normalizada.startswith('---'):
            continue
        
        # Main headers (## TITLE) - Apply selected color, minimal spacing
        if linea_normalizada.startswith('## '):
            texto = linea_normalizada.replace('## ', '')
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(texto.upper())
            run.bold = True
            run.font.size = Pt(13)
            run.font.color.rgb = obtener_color_titulo(color)
            _set_run_calibri(run)
            continue
        
        # Sub headers (### subtitle) - Apply selected color, minimal spacing
        if linea_normalizada.startswith('### '):
            texto = linea_normalizada.replace('### ', '')
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(texto)
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = obtener_color_titulo(color)
            _set_run_calibri(run)
            continue
        
        # Single # header - Apply selected color
        if linea_normalizada.startswith('# ') and not linea_normalizada.startswith('## '):
            texto = linea_normalizada.replace('# ', '')
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(texto.upper())
            run.bold = True
            run.font.size = Pt(16)
            run.font.color.rgb = obtener_color_titulo(color)
            _set_run_calibri(run)
            continue
        
        # Regular paragraph - handle bold with **text**
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        
        # Check if it's a question number (starts with number and period) - add small space before
        if re.match(r'^\d+\.', linea_normalizada):
            p.paragraph_format.space_before = Pt(6)
        
        # Check if it's an answer option (starts with a), b), c), d)) - indent slightly
        if re.match(r'^[a-d]\)', linea_normalizada):
            p.paragraph_format.left_indent = Inches(0.3)
        
        agregar_texto_con_negrita(p, linea_normalizada)
    
    doc.save(path_salida)
    print(f"✅ Examen DOCX guardado: {path_salida}")
    return path_salida


def guardar_examen_como_pdf(contenido: str, path_pdf: str = "/tmp/examen.pdf", color: str = "azul elegante") -> str:
    """
    Save exam content to a formal PDF document.
    First creates DOCX, then converts to PDF.
    
    Args:
        contenido: Markdown-like text from ChatGPT
        path_pdf: Output path for the PDF file
        color: Color scheme for document styling
        
    Returns:
        Path to the saved PDF file
    """
    # First create DOCX with color
    path_docx = path_pdf.replace('.pdf', '.docx')
    guardar_examen_como_docx(contenido, path_docx, color=color)
    
    # Convert to PDF
    result_pdf = convert_to_pdf(path_docx, color=color)
    
    if result_pdf and os.path.exists(result_pdf):
        print(f"✅ Examen PDF guardado: {result_pdf}")
        return result_pdf
    
    # Fallback: use direct PDF generation
    print("⚠️ Usando fallback para PDF de examen")
    return fallback_pdf_conversion(path_docx, path_pdf, color=color)
